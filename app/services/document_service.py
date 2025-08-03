"""Document processing service for the RAG system."""

import os
import uuid
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import hashlib
import mimetypes

import chardet
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader, Docx2txtLoader
from langchain.schema import Document

from app.config import get_settings
from app.utils.exceptions import DocumentProcessingError, ValidationError
from app.utils.logger import get_logger, log_execution_time, LoggerMixin


class DocumentChunk:
    """Represents a chunk of processed document text."""
    
    def __init__(
        self,
        id: str,
        document_id: str,
        content: str,
        start_index: int,
        end_index: int,
        metadata: Optional[Dict[str, Any]] = None,
    ):
        """Initialize document chunk.
        
        Args:
            id: Unique chunk identifier
            document_id: Parent document identifier
            content: Chunk text content
            start_index: Start position in original document
            end_index: End position in original document
            metadata: Additional chunk metadata
        """
        self.id = id
        self.document_id = document_id
        self.content = content
        self.start_index = start_index
        self.end_index = end_index
        self.metadata = metadata or {}
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert chunk to dictionary."""
        return {
            "id": self.id,
            "document_id": self.document_id,
            "content": self.content,
            "start_index": self.start_index,
            "end_index": self.end_index,
            "metadata": self.metadata,
        }


class ProcessedDocument:
    """Represents a processed document with chunks and metadata."""
    
    def __init__(
        self,
        id: str,
        filename: str,
        content_type: str,
        size: int,
        chunks: List[DocumentChunk],
        metadata: Optional[Dict[str, Any]] = None,
    ):
        """Initialize processed document.
        
        Args:
            id: Unique document identifier
            filename: Original filename
            content_type: MIME content type
            size: File size in bytes
            chunks: List of document chunks
            metadata: Additional document metadata
        """
        self.id = id
        self.filename = filename
        self.content_type = content_type
        self.size = size
        self.chunks = chunks
        self.metadata = metadata or {}
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert document to dictionary."""
        return {
            "id": self.id,
            "filename": self.filename,
            "content_type": self.content_type,
            "size": self.size,
            "chunk_count": len(self.chunks),
            "chunks": [chunk.to_dict() for chunk in self.chunks],
            "metadata": self.metadata,
        }


class DocumentProcessor(LoggerMixin):
    """Service for processing documents into chunks for RAG system."""
    
    def __init__(self):
        """Initialize document processor."""
        self.settings = get_settings()
        self.supported_extensions = set(self.settings.supported_extensions)
    
    @log_execution_time("validate_file")
    def validate_file(self, file_path: str) -> Tuple[bool, Optional[str]]:
        """Validate file for processing.
        
        Args:
            file_path: Path to file to validate
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        try:
            file_path_obj = Path(file_path)
            
            # Check if file exists
            if not file_path_obj.exists():
                return False, f"File does not exist: {file_path}"
            
            # Check if it's a file (not directory)
            if not file_path_obj.is_file():
                return False, f"Path is not a file: {file_path}"
            
            # Check file extension
            file_extension = file_path_obj.suffix.lower()
            if file_extension not in self.supported_extensions:
                return False, f"Unsupported file extension: {file_extension}. Supported: {list(self.supported_extensions)}"
            
            # Check file size
            file_size = file_path_obj.stat().st_size
            if file_size > self.settings.max_file_size:
                max_size_mb = self.settings.max_file_size / (1024 * 1024)
                return False, f"File too large: {file_size} bytes. Maximum allowed: {max_size_mb}MB"
            
            # Check if file is readable
            try:
                with open(file_path, 'rb') as f:
                    f.read(1)  # Try to read one byte
            except PermissionError:
                return False, f"Permission denied reading file: {file_path}"
            except Exception as e:
                return False, f"Error reading file: {str(e)}"
            
            return True, None
            
        except Exception as e:
            self.logger.error("file_validation_error", file_path=file_path, error=str(e))
            return False, f"Validation error: {str(e)}"
    
    @log_execution_time("extract_text_from_pdf")
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentProcessingError: If PDF extraction fails
        """
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"\\n--- Page {page_num + 1} ---\\n")
                            text_content.append(page_text)
                    except Exception as e:
                        self.logger.warning(
                            "pdf_page_extraction_error",
                            file_path=file_path,
                            page_num=page_num,
                            error=str(e)
                        )
                        continue
            
            if not text_content:
                raise DocumentProcessingError(
                    "No text content extracted from PDF",
                    filename=os.path.basename(file_path)
                )
            
            return "\\n".join(text_content)
            
        except Exception as e:
            if isinstance(e, DocumentProcessingError):
                raise
            raise DocumentProcessingError(
                f"Failed to extract text from PDF: {str(e)}",
                filename=os.path.basename(file_path)
            )
    
    @log_execution_time("extract_text_from_docx")
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentProcessingError: If DOCX extraction fails
        """
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            if not text_content:
                raise DocumentProcessingError(
                    "No text content extracted from DOCX",
                    filename=os.path.basename(file_path)
                )
            
            return "\\n".join(text_content)
            
        except Exception as e:
            if isinstance(e, DocumentProcessingError):
                raise
            raise DocumentProcessingError(
                f"Failed to extract text from DOCX: {str(e)}",
                filename=os.path.basename(file_path)
            )
    
    @log_execution_time("extract_text_from_txt")
    def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from text file with encoding detection.
        
        Args:
            file_path: Path to text file
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentProcessingError: If text extraction fails
        """
        try:
            # Detect encoding
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result.get('encoding', 'utf-8')
            
            # Read with detected encoding
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
            except UnicodeDecodeError:
                # Fallback to utf-8 with error handling
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    content = file.read()
                    self.logger.warning(
                        "encoding_fallback",
                        file_path=file_path,
                        detected_encoding=encoding
                    )
            
            if not content.strip():
                raise DocumentProcessingError(
                    "No text content found in file",
                    filename=os.path.basename(file_path)
                )
            
            return content
            
        except Exception as e:
            if isinstance(e, DocumentProcessingError):
                raise
            raise DocumentProcessingError(
                f"Failed to extract text from file: {str(e)}",
                filename=os.path.basename(file_path)
            )
    
    @log_execution_time("extract_text")
    def extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from file based on type.
        
        Args:
            file_path: Path to file
            file_type: File extension (with dot)
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentProcessingError: If extraction fails
            ValidationError: If file type is unsupported
        """
        file_type = file_type.lower()
        
        if file_type not in self.supported_extensions:
            raise ValidationError(
                f"Unsupported file type: {file_type}",
                field="file_type",
                value=file_type
            )
        
        self.logger.info("extracting_text", file_path=file_path, file_type=file_type)
        
        if file_type == '.pdf':
            return self._extract_text_from_pdf(file_path)
        elif file_type == '.docx':
            return self._extract_text_from_docx(file_path)
        elif file_type in ['.txt', '.md']:
            return self._extract_text_from_txt(file_path)
        else:
            raise ValidationError(
                f"No extraction method for file type: {file_type}",
                field="file_type",
                value=file_type
            )
    
    @log_execution_time("create_chunks")
    def create_chunks(
        self,
        text: str,
        chunk_size: Optional[int] = None,
        overlap: Optional[int] = None
    ) -> List[str]:
        """Create overlapping text chunks from text using LangChain's text splitter.
        
        Args:
            text: Text to chunk
            chunk_size: Size of each chunk (uses config default if None)
            overlap: Overlap between chunks (uses config default if None)
            
        Returns:
            List of text chunks
            
        Raises:
            ValidationError: If chunk parameters are invalid
        """
        chunk_size = chunk_size or self.settings.chunk_size
        overlap = overlap or self.settings.chunk_overlap
        
        if chunk_size <= 0:
            raise ValidationError("Chunk size must be positive", field="chunk_size", value=chunk_size)
        
        if overlap < 0:
            raise ValidationError("Overlap must be non-negative", field="overlap", value=overlap)
        
        if overlap >= chunk_size:
            raise ValidationError("Overlap must be less than chunk size", field="overlap", value=overlap)
        
        # Clean text
        text = text.strip()
        if not text:
            return []
        
        # Use LangChain's RecursiveCharacterTextSplitter for better chunking
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=overlap,
            separators=["\n\n", "\n", " ", ""],
            keep_separator=False,
            length_function=len,
        )
        
        chunks = text_splitter.split_text(text)
        
        self.logger.info(
            "chunks_created_with_langchain",
            original_length=len(text),
            chunk_count=len(chunks),
            chunk_size=chunk_size,
            overlap=overlap
        )
        
        return chunks
    
    @log_execution_time("load_document_with_langchain")
    def load_document_with_langchain(self, file_path: str, file_type: str) -> List[Document]:
        """Load document using LangChain document loaders.
        
        Args:
            file_path: Path to the document
            file_type: File extension (with dot)
            
        Returns:
            List of LangChain Document objects
            
        Raises:
            DocumentProcessingError: If loading fails
            ValidationError: If file type is unsupported
        """
        file_type = file_type.lower()
        
        try:
            if file_type == '.pdf':
                loader = PyPDFLoader(file_path)
            elif file_type == '.txt' or file_type == '.md':
                loader = TextLoader(file_path, encoding='utf-8')
            elif file_type == '.docx':
                loader = Docx2txtLoader(file_path)
            else:
                raise ValidationError(
                    f"No LangChain loader for file type: {file_type}",
                    field="file_type",
                    value=file_type
                )
            
            documents = loader.load()
            
            self.logger.info(
                "document_loaded_with_langchain",
                file_path=file_path,
                file_type=file_type,
                page_count=len(documents)
            )
            
            return documents
            
        except Exception as e:
            if isinstance(e, ValidationError):
                raise
            raise DocumentProcessingError(
                f"Failed to load document with LangChain: {str(e)}",
                filename=os.path.basename(file_path)
            )
    
    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from file.
        
        Args:
            file_path: Path to file
            
        Returns:
            File metadata
        """
        try:
            file_path_obj = Path(file_path)
            stat = file_path_obj.stat()
            
            # Get MIME type
            mime_type, _ = mimetypes.guess_type(file_path)
            
            # Calculate file hash
            hash_md5 = hashlib.md5()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            
            metadata = {
                "filename": file_path_obj.name,
                "file_extension": file_path_obj.suffix.lower(),
                "file_size": stat.st_size,
                "content_type": mime_type or "application/octet-stream",
                "created_time": stat.st_ctime,
                "modified_time": stat.st_mtime,
                "file_hash": hash_md5.hexdigest(),
            }
            
            return metadata
            
        except Exception as e:
            self.logger.warning("metadata_extraction_error", file_path=file_path, error=str(e))
            return {
                "filename": os.path.basename(file_path),
                "error": f"Metadata extraction failed: {str(e)}"
            }
    
    @log_execution_time("process_file")
    async def process_file(
        self,
        file_path: str,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None,
        additional_metadata: Optional[Dict[str, Any]] = None
    ) -> ProcessedDocument:
        """Process file into document chunks.
        
        Args:
            file_path: Path to file to process
            chunk_size: Custom chunk size
            chunk_overlap: Custom chunk overlap
            additional_metadata: Additional metadata to include
            
        Returns:
            Processed document with chunks
            
        Raises:
            DocumentProcessingError: If processing fails
            ValidationError: If validation fails
        """
        # Validate file
        is_valid, error_message = self.validate_file(file_path)
        if not is_valid:
            raise ValidationError(error_message, field="file_path", value=file_path)
        
        # Generate document ID
        document_id = str(uuid.uuid4())
        
        # Get file metadata
        file_metadata = self.get_metadata(file_path)
        if additional_metadata:
            file_metadata.update(additional_metadata)
        
        # Extract file type
        file_extension = Path(file_path).suffix.lower()
        
        self.logger.info(
            "processing_document",
            document_id=document_id,
            filename=file_metadata.get("filename"),
            file_size=file_metadata.get("file_size"),
            file_type=file_extension
        )
        
        try:
            # Extract text
            text_content = self.extract_text(file_path, file_extension)
            
            # Create chunks
            text_chunks = self.create_chunks(text_content, chunk_size, chunk_overlap)
            
            # Create DocumentChunk objects
            chunks = []
            current_position = 0
            
            for i, chunk_text in enumerate(text_chunks):
                chunk_id = f"{document_id}_chunk_{i}"
                
                # Find chunk position in original text
                chunk_start = text_content.find(chunk_text, current_position)
                if chunk_start == -1:
                    chunk_start = current_position
                chunk_end = chunk_start + len(chunk_text)
                
                chunk = DocumentChunk(
                    id=chunk_id,
                    document_id=document_id,
                    content=chunk_text,
                    start_index=chunk_start,
                    end_index=chunk_end,
                    metadata={
                        "chunk_index": i,
                        "chunk_length": len(chunk_text),
                    }
                )
                chunks.append(chunk)
                current_position = chunk_end
            
            # Create processed document
            processed_doc = ProcessedDocument(
                id=document_id,
                filename=file_metadata.get("filename", os.path.basename(file_path)),
                content_type=file_metadata.get("content_type", "application/octet-stream"),
                size=file_metadata.get("file_size", 0),
                chunks=chunks,
                metadata=file_metadata
            )
            
            self.logger.info(
                "document_processed_successfully",
                document_id=document_id,
                chunk_count=len(chunks),
                total_characters=len(text_content)
            )
            
            return processed_doc
            
        except Exception as e:
            self.logger.error(
                "document_processing_failed",
                document_id=document_id,
                filename=file_metadata.get("filename"),
                error=str(e)
            )
            
            if isinstance(e, (DocumentProcessingError, ValidationError)):
                raise
            
            raise DocumentProcessingError(
                f"Unexpected error processing document: {str(e)}",
                filename=file_metadata.get("filename")
            )